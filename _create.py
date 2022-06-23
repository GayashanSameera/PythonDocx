import docx
from pptx.util import Pt
import time

start_time = time.perf_counter ()
doc = docx.Document()
doc.add_heading('gayashan', 0)
data = (
	(1, 'gaya 1'),
	(2, 'gaya 2'),
	(3, 'gaya 3'),
    (1, 'gaya 4'),
	(2, 'gaya 5'),
	(3, 'gaya 6'),
    (1, 'gaya 7'),
	(2, 'gaya 8'),
	(3, 'gaya 9'),
    (1, 'gaya 10'),
    (1, 'gaya 11'),
	(2, 'gaya 12'),
	(3, 'gaya 13'),
    (1, 'gaya 14'),
	(2, 'gaya 15'),
	(3, 'gaya 16'),
    (1, 'gaya 17'),
	(2, 'gaya 18'),
	(3, 'gaya 19'),
    (1, 'gaya 20')
)
i = 0
while i < 1500:
    table = doc.add_table(rows=1, cols=6, style='Colorful List')
    row = table.rows[0].cells
    row[0].text = 'Id'
    row[1].text = 'Name'
    row[2].text = 'aaa'
    row[3].text = 'bbbb'
    row[4].text = 'cccc'
    row[5].text = 'dddd'


    for id, name in data:
        row = table.add_row().cells
        row[0].text = str(id)
        row[1].text = "name"
        row[2].text = "name 1"
        row[3].text = "nam  2"
        row[3].text = "name 3"
        row[4].text = name


    paragraph = doc.add_paragraph(' ')
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(5)
    i += 1


doc.save('output.docx')
end_time = time.perf_counter ()
print(end_time - start_time, "seconds")
