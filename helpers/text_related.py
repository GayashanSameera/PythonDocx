from docx import Document
import pydash

from helpers.utils import replace_tags, get_tag_content, get_tag_from_string

def add_headings(doc, text, level):
    doc.add_heading(text, level)

def add_header_parah(doc, text, level):
    # Choosing the top most section of the page
    section = doc.sections[0]
    
    # Calling the header
    header = section.header
    
    # Calling the paragraph already present in
    # the header section
    header_para = header.paragraphs[0]
    # Adding text in the header
    header_para.text = text

    # Adding the centred zoned header
    # header_para.text = "\tThis is Centred Zoned Header..."

    # Adding the right zoned header
    # header_para.text = "\t\tThis is Right Zoned Header..."

def add_footer(doc, text):
    # Choosing the top most section of the page
    section = doc.sections[0]
    
    # Calling the footer
    footer = section.footer
    
    # Calling the paragraph already present in
    # the footer section
    footer_para = footer.paragraphs[0]
    
    # Adding the left zoned footer
    footer_para.text = text

def add_page_break(doc):
    doc.add_page_break()

def change_orientation_of_section(section):
    # Selecting a section of the document
    # section = doc.sections[0]

    # Changing the orientation to landscape
    section.orientation = WD_ORIENT.LANDSCAPE #PORTRAIT

def add_paragraph(doc, text, option):
    #Adding paragraph
    para = doc.add_paragraph(text)

    if(option == "page_break_before"):
        # Setting page_break_before as True
        para.page_break_before = True

    elif(option == "keep_with_next"):
        # Setting keep_with_next as True
        para.keep_with_next = True

    elif(option == "keep_together"):
        # Setting keep_together as True
        para.keep_together = True

    elif(option == "widow_control"):
        # Setting widow_control as True
        para.widow_control = True

def replace_txt(document, paragraph, replacements):
    
    pattern = r'\<PT (.*?) \>'
    matches = get_tag_content(pattern, paragraph)
    for match in matches:
        object_value = pydash.get(replacements, match)
        replace_tags(str(f"<PT {match} >"), str(object_value), paragraph)

    for section in document.sections:
        footer = section.footer
        header = section.header
        for paragraph in footer.paragraphs:
            pattern = r'\<PT (.*?) \>'
            matches = get_tag_content(pattern, paragraph)

            for match in matches:
                object_value = pydash.get(replacements, match)
                replace_tags(str(f"<PT {match} >"), str(object_value), paragraph)

        for paragraph in header.paragraphs:
            pattern = r'\<PT (.*?) \>'
            matches = get_tag_content(pattern, paragraph)

            for match in matches:
                object_value = pydash.get(replacements, match)
                replace_tags(str(f"<PT {match} >"), str(object_value), paragraph)

    #how to replace tags inside table
    # for table in document.tables:
    #     for row in table.rows:
    #         for cell in row.cells:
    #             for paragraph in cell.paragraphs:
    #                 for match, replacement in replacements.items():
    #                     replace_text_in_paragraph(paragraph, match, replacement)
